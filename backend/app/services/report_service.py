from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from io import BytesIO
from typing import List, Dict, Any
from datetime import datetime
from sqlalchemy.orm import Session

from ..models import Job, Candidate, MatchResult


class ReportService:
    """Service to generate DOCX reports"""

    def __init__(self, db: Session):
        self.db = db

    def generate_candidate_report(self, match_result: MatchResult) -> BytesIO:
        """Generate detailed report for a single candidate"""
        doc = Document()
        candidate = match_result.candidate
        job = match_result.job

        # Title
        title = doc.add_heading('Candidate Evaluation Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Header info
        doc.add_paragraph(f"African Union Commission")
        doc.add_paragraph(f"Human Resources Management Directorate")
        doc.add_paragraph(f"Talent Acquisition Unit")
        doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
        doc.add_paragraph()

        # Job Information
        doc.add_heading('Position Information', level=1)
        job_table = doc.add_table(rows=4, cols=2)
        job_table.style = 'Table Grid'
        job_data = [
            ("Position Title", job.title),
            ("Reference Number", job.reference_number or "N/A"),
            ("Grade Level", job.grade_level.value if job.grade_level else "N/A"),
            ("Minimum Pass Mark", f"{job.min_pass_mark}%")
        ]
        for i, (label, value) in enumerate(job_data):
            job_table.rows[i].cells[0].text = label
            job_table.rows[i].cells[1].text = str(value)

        doc.add_paragraph()

        # Candidate Information
        doc.add_heading('Candidate Information', level=1)
        cand_table = doc.add_table(rows=6, cols=2)
        cand_table.style = 'Table Grid'
        cand_data = [
            ("Full Name", candidate.full_name),
            ("Email", candidate.email or "N/A"),
            ("Nationality", candidate.nationality or "N/A"),
            ("Gender", candidate.gender.value.title() if candidate.gender else "Not Specified"),
            ("Least Represented Country", "Yes" if candidate.is_least_represented_country else "No"),
            ("Disability Status", "Yes" if candidate.has_disability else "No")
        ]
        for i, (label, value) in enumerate(cand_data):
            cand_table.rows[i].cells[0].text = label
            cand_table.rows[i].cells[1].text = str(value)

        doc.add_paragraph()

        # Score Summary
        doc.add_heading('Score Summary', level=1)

        # Final score highlight
        p = doc.add_paragraph()
        run = p.add_run(f"FINAL SCORE: {match_result.final_score:.1f}/100")
        run.bold = True
        run.font.size = Pt(16)

        p = doc.add_paragraph()
        run = p.add_run(f"RANK: #{match_result.rank}")
        run.bold = True

        status = "PASSES" if match_result.passes_cutoff else "DOES NOT PASS"
        p = doc.add_paragraph()
        run = p.add_run(f"Cutoff Status: {status} (minimum: {job.min_pass_mark}%)")
        run.bold = True

        doc.add_paragraph()

        # Score breakdown table
        score_table = doc.add_table(rows=6, cols=2)
        score_table.style = 'Table Grid'
        score_data = [
            ("Education Score (30%)", f"{match_result.education_total:.1f}/30"),
            ("Experience Score (70%)", f"{match_result.experience_total:.1f}/70"),
            ("Base Score", f"{match_result.base_score:.1f}/100"),
            ("Female Bonus", f"+{match_result.bonus_female}"),
            ("Age Bonus (≤35)", f"+{match_result.bonus_age}"),
            ("Least Represented Bonus", f"+{match_result.bonus_least_represented}"),
        ]
        for i, (label, value) in enumerate(score_data):
            score_table.rows[i].cells[0].text = label
            score_table.rows[i].cells[1].text = str(value)

        doc.add_paragraph()

        # Education Scores Detail
        doc.add_heading('Education Assessment (30%)', level=1)
        if match_result.education_scores:
            for criterion_id, data in match_result.education_scores.items():
                p = doc.add_paragraph()
                run = p.add_run(f"{criterion_id.replace('_', ' ').title()}: ")
                run.bold = True
                p.add_run(f"{data.get('score', 0)}/{data.get('max', 10)}")

                reasoning = doc.add_paragraph(data.get('reasoning', 'No reasoning provided'))
                reasoning.style = 'Quote'

        doc.add_paragraph()

        # Experience Scores Detail
        doc.add_heading('Experience Assessment (70%)', level=1)
        if match_result.experience_scores:
            for criterion_id, data in match_result.experience_scores.items():
                p = doc.add_paragraph()
                run = p.add_run(f"{criterion_id.replace('_', ' ').title()}: ")
                run.bold = True
                p.add_run(f"{data.get('score', 0)}/{data.get('max', 10)}")

                reasoning = doc.add_paragraph(data.get('reasoning', 'No reasoning provided'))
                reasoning.style = 'Quote'

        doc.add_paragraph()

        # Overall Assessment
        doc.add_heading('Overall Assessment', level=1)
        doc.add_paragraph(match_result.overall_reasoning or "No overall assessment provided")

        # Strengths
        doc.add_heading('Strengths', level=2)
        if match_result.strengths:
            for strength in match_result.strengths:
                doc.add_paragraph(strength, style='List Bullet')
        else:
            doc.add_paragraph("No strengths identified")

        # Weaknesses
        doc.add_heading('Areas of Concern / Gaps', level=2)
        if match_result.weaknesses:
            for weakness in match_result.weaknesses:
                doc.add_paragraph(weakness, style='List Bullet')
        else:
            doc.add_paragraph("No concerns identified")

        # Flags
        if match_result.flags:
            doc.add_heading('Flags', level=2)
            for flag in match_result.flags:
                doc.add_paragraph(f"⚠ {flag}", style='List Bullet')

        # Recommendations
        doc.add_heading('Recommendations', level=1)
        doc.add_paragraph(match_result.recommendations or "No recommendations provided")

        # Footer
        doc.add_paragraph()
        doc.add_paragraph("─" * 50)
        p = doc.add_paragraph()
        run = p.add_run("This report was generated by the AI-supported CV Matching Tool")
        run.italic = True
        p = doc.add_paragraph()
        run = p.add_run("African Union Commission - Human Resources Management Directorate")
        run.italic = True

        # Save to BytesIO
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        return buffer

    def generate_longlist_report(self, job_id: int) -> BytesIO:
        """Generate longlist report for a job (top 20 candidates)"""
        job = self.db.query(Job).filter(Job.id == job_id).first()
        if not job:
            raise ValueError("Job not found")

        results = self.db.query(MatchResult).filter(
            MatchResult.job_id == job_id
        ).order_by(MatchResult.final_score.desc()).limit(20).all()

        doc = Document()

        # Title
        title = doc.add_heading('Candidate Longlist Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Header
        doc.add_paragraph("African Union Commission")
        doc.add_paragraph("Human Resources Management Directorate")
        doc.add_paragraph("Talent Acquisition Unit")
        doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
        doc.add_paragraph()

        # Job Info
        doc.add_heading('Position Details', level=1)
        doc.add_paragraph(f"Position: {job.title}")
        doc.add_paragraph(f"Reference: {job.reference_number or 'N/A'}")
        doc.add_paragraph(f"Grade: {job.grade_level.value if job.grade_level else 'N/A'}")
        doc.add_paragraph(f"Minimum Pass Mark: {job.min_pass_mark}%")
        doc.add_paragraph()

        # Summary Statistics
        doc.add_heading('Summary Statistics', level=1)
        total_candidates = self.db.query(Candidate).filter(
            Candidate.job_id == job_id
        ).count()

        passing = sum(1 for r in results if r.passes_cutoff)
        doc.add_paragraph(f"Total Candidates Screened: {total_candidates}")
        doc.add_paragraph(f"Candidates in Longlist: {len(results)}")
        doc.add_paragraph(f"Meeting Cutoff Score: {passing}")

        doc.add_paragraph()

        # Longlist Table
        doc.add_heading('Top 20 Candidates (Longlist)', level=1)

        table = doc.add_table(rows=1, cols=7)
        table.style = 'Table Grid'

        # Header row
        headers = ['Rank', 'Name', 'Gender', 'Nationality', 'Education', 'Experience', 'Final Score']
        header_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            header_cells[i].text = header
            header_cells[i].paragraphs[0].runs[0].bold = True

        # Data rows
        for result in results:
            candidate = result.candidate
            row = table.add_row().cells
            row[0].text = str(result.rank)
            row[1].text = candidate.full_name
            row[2].text = candidate.gender.value.title() if candidate.gender else "N/S"
            row[3].text = candidate.nationality or "N/A"
            row[4].text = f"{result.education_total:.1f}/30"
            row[5].text = f"{result.experience_total:.1f}/70"
            row[6].text = f"{result.final_score:.1f}"

        doc.add_paragraph()

        # Individual summaries
        doc.add_heading('Individual Candidate Summaries', level=1)

        for result in results:
            candidate = result.candidate
            doc.add_heading(f"#{result.rank}. {candidate.full_name}", level=2)

            p = doc.add_paragraph()
            run = p.add_run(f"Score: {result.final_score:.1f}/100")
            run.bold = True

            if result.passes_cutoff:
                p.add_run(" ✓ Meets cutoff")
            else:
                p.add_run(" ✗ Below cutoff")

            # Bonus breakdown
            bonuses = []
            if result.bonus_female > 0:
                bonuses.append("Female (+5)")
            if result.bonus_age > 0:
                bonuses.append("Age ≤35 (+5)")
            if result.bonus_least_represented > 0:
                bonuses.append("Least Rep. Country (+5)")
            if result.bonus_inclusion > 0:
                bonuses.append("Inclusion (+5)")

            if bonuses:
                doc.add_paragraph(f"Bonuses: {', '.join(bonuses)}")

            doc.add_paragraph(f"Summary: {result.overall_reasoning[:500] if result.overall_reasoning else 'N/A'}...")
            doc.add_paragraph()

        # Footer
        doc.add_paragraph("─" * 50)
        p = doc.add_paragraph()
        run = p.add_run("Generated by AI-supported CV Matching Tool - African Union Commission")
        run.italic = True

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        return buffer

    def generate_excel_report(self, job_id: int) -> BytesIO:
        """Generate Excel report with all candidates"""
        import openpyxl
        from openpyxl.styles import Font, PatternFill, Alignment

        job = self.db.query(Job).filter(Job.id == job_id).first()
        results = self.db.query(MatchResult).filter(
            MatchResult.job_id == job_id
        ).order_by(MatchResult.final_score.desc()).all()

        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "Candidate Rankings"

        # Headers
        headers = [
            'Rank', 'Full Name', 'Email', 'Gender', 'Nationality',
            'Education Score', 'Experience Score', 'Base Score',
            'Female Bonus', 'Age Bonus', 'LRC Bonus', 'Inclusion Bonus',
            'Total Bonus', 'Final Score', 'Passes Cutoff', 'In Longlist'
        ]

        for col, header in enumerate(headers, 1):
            cell = ws.cell(row=1, column=col, value=header)
            cell.font = Font(bold=True)
            cell.fill = PatternFill(start_color="366092", end_color="366092", fill_type="solid")
            cell.font = Font(bold=True, color="FFFFFF")

        # Data
        for row, result in enumerate(results, 2):
            candidate = result.candidate
            ws.cell(row=row, column=1, value=result.rank)
            ws.cell(row=row, column=2, value=candidate.full_name)
            ws.cell(row=row, column=3, value=candidate.email or "")
            ws.cell(row=row, column=4, value=candidate.gender.value if candidate.gender else "")
            ws.cell(row=row, column=5, value=candidate.nationality or "")
            ws.cell(row=row, column=6, value=result.education_total)
            ws.cell(row=row, column=7, value=result.experience_total)
            ws.cell(row=row, column=8, value=result.base_score)
            ws.cell(row=row, column=9, value=result.bonus_female)
            ws.cell(row=row, column=10, value=result.bonus_age)
            ws.cell(row=row, column=11, value=result.bonus_least_represented)
            ws.cell(row=row, column=12, value=result.bonus_inclusion)
            ws.cell(row=row, column=13, value=result.total_bonus)
            ws.cell(row=row, column=14, value=result.final_score)
            ws.cell(row=row, column=15, value="Yes" if result.passes_cutoff else "No")
            ws.cell(row=row, column=16, value="Yes" if result.is_in_longlist else "No")

        # Auto-adjust column widths
        for column in ws.columns:
            max_length = 0
            column_letter = column[0].column_letter
            for cell in column:
                try:
                    if len(str(cell.value)) > max_length:
                        max_length = len(str(cell.value))
                except:
                    pass
            adjusted_width = (max_length + 2)
            ws.column_dimensions[column_letter].width = adjusted_width

        buffer = BytesIO()
        wb.save(buffer)
        buffer.seek(0)

        return buffer
